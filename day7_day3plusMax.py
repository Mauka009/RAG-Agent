# ============================================================
# 文件名: day7_day3plusMax.py
# 说明: 改进后的RAG问答系统，包含：
#       - 向量库本地加载/创建
#       - BM25关键词检索
#       - 混合检索 + Cross-Encoder重排序
#       - 多轮对话
#       - 完善错误处理
# ============================================================

from dotenv import load_dotenv
load_dotenv()                       # 加载 .env 中的环境变量

import config

import os
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"
from pypdf import PdfReader
from PIL import Image
import re
import json
import requests
import numpy as np
import gradio as gr
import jieba
import pytesseract
import time
from PIL import Image
from rank_bm25 import BM25Okapi
from sentence_transformers import CrossEncoder

# LangChain 相关
from langchain_community.document_loaders import TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document   # 用于构造临时文档
from langchain_community.document_loaders import UnstructuredWordDocumentLoader, UnstructuredExcelLoader

# ---------- 1. 配置与常量 ---------

# 文件路径
TEXT_FILE_PATH = config.TEXT_FILE_PATH
INDEX_PATH = config.INDEX_PATH

# 模型名称
EMBEDDING_MODEL = config.EMBEDDING_MODEL
RERANK_MODEL = config.RERANK_MODEL

# 混合检索权重
WEIGHT_BM25 = config.WEIGHT_BM25
WEIGHT_VECTOR = config.WEIGHT_VECTOR

# ---------- 2. 检查环境变量 ----------
API_KEY = config.API_KEY
if not API_KEY:
    raise ValueError("❌ 请在 .env 文件中设置 SILICONFLOW_API_KEY")
BASE_URL = config.BASE_URL

# ---------- 2.5 智能文档加载器 ----------

def smart_loader(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.txt':
        from langchain_community.document_loaders import TextLoader
        loader = TextLoader(file_path, encoding='utf-8')
        return loader.load()
    
    elif ext == '.pdf':
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        if not text.strip():
            print("⚠️ 警告:该PDF未提取到文字,可能是扫描件")
        return [Document(page_content=text, metadata={"source": file_path})]
    
    elif ext in ['.jpg', '.jpeg', '.png', '.bmp']:
        # 用 pytesseract 识别图片文字
        img = Image.open(file_path)
        # 中文 + 英文
        text = pytesseract.image_to_string(img, lang='chi_sim+eng')
        if not text.strip():
            print(f"⚠️ 图片 {file_path} 未识别到文字")
        return [Document(page_content=text.strip(), metadata={"source": file_path})]
    
    elif ext in ['.docx','.DOCX']:
        try:
            # 优先尝试 Unstructured（保留标题等结构）
            from langchain_community.document_loaders import UnstructuredWordDocumentLoader
            loader = UnstructuredWordDocumentLoader(file_path, mode="elements")
            docs = loader.load()
            if docs:
                print("✅ 使用 Unstructured 解析 docx")
                return docs
        except Exception as e:
            print(f"⚠️ Unstructured 解析失败，降级到 python-docx: {e}")
    
    # 降级方案：用 python-docx 兜底
        import docx
        doc = docx.Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return [Document(page_content=text, metadata={"source": file_path})]
    else:
        raise ValueError(f"❌ 不支持的文件格式: {ext}")

# ---------- 3. 初始化 Embedding 模型 ----------
embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)

# ---------- 4. 加载或创建向量库 ----------
def load_or_create_vectorstore():
    # 如果索引存在 且 不强制重建，才加载缓存
    if os.path.exists(INDEX_PATH) and not config.FORCE_REBUILD:
        print("📂 发现本地向量库，正在加载...")
        return FAISS.load_local(INDEX_PATH, embeddings, allow_dangerous_deserialization=True)
    
    print("📄 未找到向量库，从文档创建...")
    if not os.path.exists(TEXT_FILE_PATH):
        raise FileNotFoundError(f"文档文件不存在: {TEXT_FILE_PATH}")
    
    # ⚠️ 核心改动：判断是文件还是文件夹
    documents = []
    if os.path.isfile(TEXT_FILE_PATH):
        # ---- 模式1：单文件（你原来的方式） ----
        print(f"📄 检测到单文件模式：{TEXT_FILE_PATH}")
        documents = smart_loader(TEXT_FILE_PATH)
    else:
        # ---- 模式2：多文件文件夹（新增功能） ----
        print(f"📁 检测到文件夹模式：{TEXT_FILE_PATH}")
        supported_exts = ('.txt', '.pdf', '.jpg', '.jpeg', '.png', '.bmp','.docx')
        for filename in os.listdir(TEXT_FILE_PATH):
            if filename.lower().endswith(supported_exts):
                file_path = os.path.join(TEXT_FILE_PATH, filename)
                try:
                    docs = smart_loader(file_path)
                    documents.extend(docs)
                    print(f"  ✅ 已加载: {filename}")
                except Exception as e:
                    print(f"  ❌ 加载失败 {filename}: {e}")
        print(f"📚 共从文件夹加载了 {len(documents)} 个文档块")
    
    # ⚠️ 如果加载到的文档是空的，报错退出
    if not documents:
        raise ValueError("❌ 未能加载到任何有效文档，请检查文件格式或路径")
    
    # ---- 切片 ----
    CHUNK_SIZE = config.CHUNK_SIZE
    CHUNK_OVERLAP = config.CHUNK_OVERLAP
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    chunks = text_splitter.split_documents(documents)
    print(f"✂️ 切分成 {len(chunks)} 个文本块")
    
    # ---- 创建向量库 ----
    vectorstore = FAISS.from_documents(chunks, embeddings)
    vectorstore.save_local(INDEX_PATH)
    print("✅ 向量库已保存到本地")
    
    # ---- 返回向量库 ----
    return vectorstore
vectorstore = load_or_create_vectorstore()

# ---------- 5. 提取所有文本块（用于BM25） ----------
# 注意：这里使用 vectorstore.docstore._dict 是 LangChain 的内部实现，
# 但为了通用性，我们也可以用遍历的方式，但目前的版本稳定可用。
# 如果你担心，可以改用 vectorstore.similarity_search("", k=10000) 但效率低。
all_texts = []
for doc_id in vectorstore.docstore._dict:
    doc = vectorstore.docstore._dict[doc_id]
    all_texts.append(doc.page_content)
print(f"📚 共 {len(all_texts)} 个文本块用于BM25")

# ---------- 6. 构建BM25检索器 ----------
def build_bm25(texts):
    tokenized_corpus = []
    for text in texts:
        # 去标点
        clean_text = re.sub(r'[^\w\u4e00-\u9fff]', ' ', text)
        words = jieba.lcut(clean_text)
        words = [w.strip() for w in words if w.strip()]
        tokenized_corpus.append(words)
    return BM25Okapi(tokenized_corpus)

bm25 = build_bm25(all_texts)
print("✅ BM25检索器构建完成")

# ---------- 7. 加载Rerank模型 ----------
print("⏳ 正在加载Rerank模型(首次会下载）...")
reranker = CrossEncoder(RERANK_MODEL)
print("✅ Rerank模型加载完成")

# ---------- 8. 核心检索函数：混合检索 + Rerank ----------
def hybrid_search_with_rerank(query, top_k=3, use_bm25=True, use_rerank=True, alpha=0.7):
    """
    输入问题，返回最相关的 top_k 个文档块(Document 对象列表）
    - use_bm25: True 使用 BM25 混合；False 只用 FAISS
    - use_rerank: True 使用 Rerank 精排；False 直接按融合分数排序取 top_k
    - alpha: 向量分数的权重，BM25 权重为 1-alpha（仅在 use_bm25=True 且 use_rerank=False 时生效）
    """
    # 8.1 向量检索（带分数）
    vector_results = vectorstore.similarity_search_with_score(query, k=top_k * 3)
    vector_docs = [doc for doc, _ in vector_results]
    vector_scores = {doc.page_content: score for doc, score in vector_results}
    
    # 构建候选字典（保留向量检索的所有文档）
    candidates_dict = {doc.page_content: doc for doc in vector_docs}
    
    # 如果使用 BM25
    if use_bm25:
        query_words = jieba.lcut(re.sub(r'[^\w\u4e00-\u9fff]', ' ', query))
        query_words = [w.strip() for w in query_words if w.strip()]
        bm25_scores = bm25.get_scores(query_words)
        # 获取 BM25 top 文档
        bm25_indices = sorted(range(len(bm25_scores)), key=lambda i: bm25_scores[i], reverse=True)[:top_k * 3]
        bm25_docs = [Document(page_content=all_texts[i]) for i in bm25_indices]
        bm25_score_dict = {all_texts[i]: bm25_scores[i] for i in bm25_indices}
        # 加入候选集
        for text in bm25_score_dict:
            if text not in candidates_dict:
                candidates_dict[text] = Document(page_content=text)
    else:
        bm25_score_dict = {}
    
    # 如果不用 Rerank，则根据融合分数排序
    if not use_rerank:
        # 准备所有文档内容
        all_contents = list(candidates_dict.keys())
        if not all_contents:
            return []
        
        # 归一化向量分数（min-max）
        vec_vals = [vector_scores.get(c, 0) for c in all_contents]
        vec_min, vec_max = min(vec_vals), max(vec_vals)
        if vec_max == vec_min:
            vec_norm = {c: 0.5 for c in all_contents}
        else:
            vec_norm = {c: (vector_scores.get(c, 0) - vec_min) / (vec_max - vec_min) for c in all_contents}
        
        if use_bm25:
            bm_vals = [bm25_score_dict.get(c, 0) for c in all_contents]
            bm_min, bm_max = min(bm_vals), max(bm_vals)
            if bm_max == bm_min:
                bm_norm = {c: 0.5 for c in all_contents}
            else:
                bm_norm = {c: (bm25_score_dict.get(c, 0) - bm_min) / (bm_max - bm_min) for c in all_contents}
        else:
            bm_norm = {c: 0 for c in all_contents}
        
        # 融合分数
        combined = []
        for content in all_contents:
            score = alpha * vec_norm[content] + (1 - alpha) * bm_norm[content]
            combined.append((candidates_dict[content], score))
        combined.sort(key=lambda x: x[1], reverse=True)
        return [doc for doc, _ in combined[:top_k]]
    
    # 否则使用 Rerank（原逻辑）
    docs = list(candidates_dict.values())
    if not docs:
        return []
    pairs = [[query, doc.page_content] for doc in docs]
    scores = reranker.predict(pairs)
    doc_score_pairs = list(zip(docs, scores))
    doc_score_pairs.sort(key=lambda x: x[1], reverse=True)
    return [doc for doc, _ in doc_score_pairs[:top_k]]
# ---------- 9. 调用大模型问答 ----------
def ask_rag(question, history, use_bm25=True, use_rerank=True):
    """
    Gradio ChatInterface 要求的函数签名: fn(question, history)
    """
    # ---------- 打印当前处理的问题 ----------
    print(f"⏳ 正在检索: {question[:30]}...")
    
    # 9.1 检索相关文档
    retrieved_docs = hybrid_search_with_rerank(question, top_k=3, use_bm25=use_bm25, use_rerank=use_rerank)
    if not retrieved_docs:
        context = "未找到相关资料"
    else:
        context = "\n\n".join([doc.page_content for doc in retrieved_docs])
    
    # ---------- 打印检索完成 ----------
    print(f"✅ 检索完成，共 {len(retrieved_docs)} 个文档块")
    
    # 9.2 构造 system prompt
    system_prompt = f"""你是一个基于知识库回答问题的AI助手。
请根据以下参考资料回答用户的问题。如果资料里没有相关信息，请直接说"资料中没有提到"。

### 参考资料：
{context}
"""
    # 构建消息列表
    messages = [{"role": "system", "content": system_prompt}]
    
    for turn in history:
        if isinstance(turn, dict):
            if "role" in turn and "content" in turn:
                messages.append({"role": turn["role"], "content": turn["content"]})
            elif "user" in turn and "bot" in turn:
                messages.append({"role": "user", "content": turn["user"]})
                messages.append({"role": "assistant", "content": turn["bot"]})
        elif isinstance(turn, (list, tuple)):
            if len(turn) >= 2:
                messages.append({"role": "user", "content": turn[0]})
                messages.append({"role": "assistant", "content": turn[1]})
            elif len(turn) == 1:
                messages.append({"role": "user", "content": turn[0]})
        else:
            print(f"⚠️ 未知的历史格式，已跳过: {turn}")
    
    messages.append({"role": "user", "content": question})
    
    # 9.4 调用 API（带重试）
    url = f"{BASE_URL}/chat/completions"
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "deepseek-ai/DeepSeek-V3",
        "messages": messages,
        "temperature": 0.7
    }
    
    print(f"⏳ 正在调用 API...")
    
    max_retries = 3
    retry_delay = 2
    
    for attempt in range(max_retries):
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=60)
            if resp.status_code == 200:
                result = resp.json()
                return result["choices"][0]["message"]["content"]
            else:
                print(f"⚠️ API 返回错误 {resp.status_code}，尝试 {attempt+1}/{max_retries}")
                time.sleep(retry_delay)
        except requests.exceptions.Timeout:
            print(f"⏰ API 请求超时，尝试 {attempt+1}/{max_retries}")
            time.sleep(retry_delay)
        except requests.exceptions.ConnectionError:
            print(f"❌ 网络连接失败，尝试 {attempt+1}/{max_retries}")
            time.sleep(retry_delay)
        except Exception as e:
            print(f"❌ 未知异常: {e}，尝试 {attempt+1}/{max_retries}")
            time.sleep(retry_delay)
    
    return "❌ API 调用失败，重试多次后仍然失败"

# ---------- 10. 启动 Gradio 界面 ----------
demo = gr.ChatInterface(
    fn=ask_rag,
    title="📚 RAG 知识库问答系统（混合检索+Rerank)",
    description="基于成都市人才引进文档回答问题，支持多轮对话。",
)

if __name__ == "__main__":
    # 启动服务，share=True 生成公网链接（可选）
    demo.launch(share=True, server_port=7860)
    print("🌐 服务已启动，访问 http://127.0.0.1:7860")